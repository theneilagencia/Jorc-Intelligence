"""
QIVO Intelligence Layer - Preprocessor Module
Extrai e limpa texto de documentos (PDF, DOCX, TXT)
"""

import os
import re
from typing import Optional, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
import aiofiles


class DocumentPreprocessor:
    """Preprocessa documentos técnicos para análise de compliance"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self):
        self.metadata: Dict[str, Any] = {}
    
    async def preprocess_text(self, file_path: str) -> str:
        """
        Extrai e limpa texto de arquivo
        
        Args:
            file_path: Caminho do arquivo
            
        Returns:
            Texto limpo e preprocessado
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Formato não suportado: {extension}")
        
        # Detectar tipo e extrair
        if extension == '.pdf':
            text = await self._extract_pdf(file_path)
        elif extension in {'.docx', '.doc'}:
            text = await self._extract_docx(file_path)
        elif extension == '.txt':
            text = await self._extract_txt(file_path)
        else:
            text = ""
        
        # Limpar e normalizar
        cleaned_text = self._clean_text(text)
        
        # Atualizar metadata
        self.metadata = {
            'file_name': path.name,
            'file_type': extension,
            'file_size': path.stat().st_size,
            'char_count': len(cleaned_text),
            'word_count': len(cleaned_text.split())
        }
        
        return cleaned_text
    
    async def _extract_pdf(self, file_path: str) -> str:
        """Extrai texto de PDF"""
        text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
        except Exception as e:
            raise ValueError(f"Erro ao ler PDF: {str(e)}")
        
        return '\n'.join(text)
    
    async def _extract_docx(self, file_path: str) -> str:
        """Extrai texto de DOCX"""
        text = []
        
        try:
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extrair texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text.append(row_text)
        except Exception as e:
            raise ValueError(f"Erro ao ler DOCX: {str(e)}")
        
        return '\n'.join(text)
    
    async def _extract_txt(self, file_path: str) -> str:
        """Extrai texto de TXT"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = await file.read()
                return text
        except Exception as e:
            raise ValueError(f"Erro ao ler TXT: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Remove ruído e normaliza texto
        
        Args:
            text: Texto bruto
            
        Returns:
            Texto limpo
        """
        # Remover múltiplos espaços
        text = re.sub(r'\s+', ' ', text)
        
        # Remover caracteres especiais excessivos
        text = re.sub(r'[^\w\s\.,;:!?()\-\[\]{}]', '', text)
        
        # Normalizar quebras de linha
        text = re.sub(r'\n+', '\n', text)
        
        # Remover espaços no início/fim
        text = text.strip()
        
        return text
    
    def get_metadata(self) -> Dict[str, Any]:
        """Retorna metadata do último documento processado"""
        return self.metadata
